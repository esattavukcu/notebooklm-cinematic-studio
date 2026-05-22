"""
bulk_import.py — Drive klasöründen toplu script import.

Public-share edilmiş Drive klasöründeki .docx dosyalarını gdown ile indirir,
python-docx ile düz metin çıkarır ve her birini Job olarak queue'ya ekler.

Kullanım (app.py'dan):
    from bulk_import import bulk_import_from_drive
    result = bulk_import_from_drive(
        drive_url_or_id="https://drive.google.com/drive/folders/XXXX",
        custom_prompt_template="Cinematic educational ...",
        submitted_by="Mehmet",
        on_progress=lambda msg: print(msg),
    )
    # result: {"created": [job_id, ...], "skipped": [...], "errors": [...]}

Drive klasörü 'Anyone with the link' olmalı — gdown API key'siz erişir.
Private klasör için OAuth Drive scope gerekir (out of scope, future iter).
"""

from __future__ import annotations

import io
import re
import shutil
import tempfile
import time
import uuid
from dataclasses import asdict
from pathlib import Path
from typing import Callable, Optional

try:
    import gdown  # type: ignore
    _GDOWN_AVAILABLE = True
except ImportError as _gd_err:
    gdown = None  # type: ignore
    _GDOWN_AVAILABLE = False
    _GDOWN_IMP_ERR = str(_gd_err)

try:
    from docx import Document as _DocxDocument  # python-docx
    _DOCX_AVAILABLE = True
except ImportError as _doc_err:
    _DocxDocument = None  # type: ignore
    _DOCX_AVAILABLE = False
    _DOCX_IMP_ERR = str(_doc_err)


# ---------------------------------------------------------------------------
# Drive URL/ID parsing
# ---------------------------------------------------------------------------
_DRIVE_FOLDER_PATTERNS = [
    # Standart: /drive/folders/<id>
    re.compile(r"drive\.google\.com/drive/folders/([a-zA-Z0-9_-]+)"),
    # Çok hesaplı oturum: /drive/u/0/folders/<id> veya /drive/u/1/folders/<id>
    re.compile(r"drive\.google\.com/drive/u/\d+/folders/([a-zA-Z0-9_-]+)"),
    # Eski folderview formatı
    re.compile(r"drive\.google\.com/folderview\?id=([a-zA-Z0-9_-]+)"),
    # ?id= veya &id= parametreli her Drive URL'i
    re.compile(r"drive\.google\.com/.*[?&]id=([a-zA-Z0-9_-]+)"),
    # Düz folder ID (20+ karakter)
    re.compile(r"^([a-zA-Z0-9_-]{20,})$"),
]


def extract_folder_id(url_or_id: str) -> Optional[str]:
    """Drive klasör URL'i veya ID'sinden plain folder ID çıkar."""
    s = (url_or_id or "").strip()
    if not s:
        return None
    for pat in _DRIVE_FOLDER_PATTERNS:
        m = pat.search(s)
        if m:
            return m.group(1)
    return None


def is_available() -> tuple[bool, str]:
    """gdown + python-docx yüklü mü?"""
    msgs = []
    if not _GDOWN_AVAILABLE:
        msgs.append(f"gdown yok: {_GDOWN_IMP_ERR}")
    if not _DOCX_AVAILABLE:
        msgs.append(f"python-docx yok: {_DOCX_IMP_ERR}")
    if msgs:
        return False, " · ".join(msgs)
    try:
        import gdown as _g  # noqa
        import docx as _d  # noqa
        return True, f"gdown v{_g.__version__} · python-docx v{_d.__version__}"
    except Exception:
        return True, "available"


# ---------------------------------------------------------------------------
# Doküman parsing — .docx, .txt, .md destekler
# ---------------------------------------------------------------------------
# Desteklenen dosya uzantıları (lowercase, dot ile). Yeni format eklemek
# için sadece bu listeyi + parse_document switch'ini güncelle.
SUPPORTED_EXTENSIONS: tuple[str, ...] = (".docx", ".txt", ".md")


def parse_docx(path: Path) -> str:
    """docx → düz metin. Paragraf ayraçları korunur, boş satırlar bir kez."""
    if not _DOCX_AVAILABLE:
        raise RuntimeError(f"python-docx yüklü değil: {_DOCX_IMP_ERR}")
    doc = _DocxDocument(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        parts.append(t)
    text = "\n".join(parts)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def parse_txt(path: Path) -> str:
    """Plain text dosya oku — UTF-8 öncelik, latin-1 fallback (Windows export)."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1254"):
        try:
            text = path.read_text(encoding=encoding)
            # Çoklu boş satırları teke indir (paragraf normalizasyonu)
            text = re.sub(r"\n{3,}", "\n\n", text).strip()
            return text
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"{path.name} text encoding tespit edilemedi (utf-8/latin-1/cp1254 hep fail)")


def parse_document(path: Path) -> str:
    """Generic parser — uzantıya göre uygun handler'a yönlendirir.

    Desteklenen: .docx (python-docx), .txt + .md (plain text).
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        return parse_docx(path)
    if ext in (".txt", ".md"):
        return parse_txt(path)
    raise RuntimeError(f"Desteklenmeyen format: {ext} ({path.name})")


def get_docx_metadata(path: Path) -> dict:
    """Doküman metadata → dict.

    .docx için: core_properties (created, modified, author, ...).
    .txt/.md için: filesystem stat (mtime → modified, sadece). docx'taki
    'author' yok, bu yüzden author='' döner.

    Döner: {created, modified, author, last_modified_by, n_paragraphs}
    'created' ve 'modified' ISO format string.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        if not _DOCX_AVAILABLE:
            return {}
        try:
            doc = _DocxDocument(str(path))
            props = doc.core_properties
            return {
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "author": (props.author or "").strip(),
                "last_modified_by": (props.last_modified_by or "").strip(),
                "n_paragraphs": len(doc.paragraphs),
            }
        except Exception:
            return {}
    # .txt / .md → sadece filesystem mtime + line count
    try:
        from datetime import datetime as _dt
        stat = path.stat()
        # Line sayısı paragraf yaklaşığı
        try:
            n_lines = sum(1 for _ in path.open("r", encoding="utf-8", errors="ignore"))
        except Exception:
            n_lines = 0
        return {
            "created": _dt.fromtimestamp(stat.st_ctime).isoformat(),
            "modified": _dt.fromtimestamp(stat.st_mtime).isoformat(),
            "author": "",
            "last_modified_by": "",
            "n_paragraphs": n_lines,
        }
    except Exception:
        return {}


def docx_to_title(filename: str, fallback_text: str = "") -> str:
    """Filename'den anlamlı title üret. '01_passion_fruit.docx' → 'Passion Fruit'."""
    stem = Path(filename).stem
    # Leading sayı + underscore/dash temizle ('01_', '01-')
    stem = re.sub(r"^\d+[_\-\.\s]+", "", stem)
    # underscore/dash → boşluk
    stem = stem.replace("_", " ").replace("-", " ").strip()
    if stem:
        return stem[:80].title()
    # Fallback: text'in ilk satırı
    if fallback_text:
        first_line = fallback_text.strip().split("\n", 1)[0][:80].strip()
        if first_line:
            return first_line
    return "Untitled"


# ---------------------------------------------------------------------------
# Drive download
# ---------------------------------------------------------------------------
def download_drive_folder(folder_id_or_url: str,
                          out_dir: Path,
                          quiet: bool = True) -> list[Path]:
    """Public Drive klasöründeki tüm dosyaları out_dir'a indir, path listesi döner.

    Birincil yol: gdown.download_folder(). Bu yöntem native Google Docs/Sheets/Slides
    içeren klasörlerde FileURLRetrievalError fırlatıp tüm indirmeyi iptal eder.

    Fallback: skip_download=True ile dosya listesi alınır, her dosya tek tek indirilir.
    Native Google Docs (uzantısız isimler) için export URL denenir; yine de başarısız
    olursa o dosya sessizce atlanır — geri kalan .txt/.docx dosyaları indirilir.
    """
    if not _GDOWN_AVAILABLE:
        raise RuntimeError(f"gdown yüklü değil: {_GDOWN_IMP_ERR}")
    folder_id = extract_folder_id(folder_id_or_url)
    if not folder_id:
        raise ValueError(
            "Geçerli Drive klasör URL'i/ID'si değil. "
            "Örnek: https://drive.google.com/drive/folders/ABCdef123..."
        )
    out_dir.mkdir(parents=True, exist_ok=True)

    # gdown 6.0.0'da ``drive.usercontent.google.com`` redirect parse bug'ı var:
    # 303 redirect sonrası "Cannot retrieve the public link" hatası fırlatıp
    # tüm klasör indirmeyi iptal ediyor. Bypass: her dosyayı tek tek requests
    # ile indir (curl ile sorunsuz çalıştığı doğrulandı). gdown sadece klasör
    # listesi almak için kullanılır (skip_download=True).
    return _download_drive_folder_per_file(
        folder_id=folder_id,
        out_dir=out_dir,
        quiet=quiet,
        first_exc=None,
    )


def _download_drive_folder_per_file(
    folder_id: str,
    out_dir: Path,
    quiet: bool = True,
    first_exc: Optional[Exception] = None,
) -> list[Path]:
    """gdown fallback: dosya listesini al, her dosyayı tek tek indir.

    gdown 6.x'in ``download_folder(skip_download=True)`` özelliği kullanılarak
    dosya listesi alınır. Her dosya ayrı ``gdown.download()`` çağrısıyla indirilir.
    Native Google Docs (klasör listesinde uzantısız isimler) için Google'ın
    export API'si denenir; yine başarısız olursa dosya atlanır.
    """
    import sys as _sys

    url = f"https://drive.google.com/drive/folders/{folder_id}"

    # skip_download=True: indirmeden sadece dosya listesini döndür
    try:
        file_list = gdown.download_folder(
            url=url,
            output=str(out_dir),
            quiet=quiet,
            use_cookies=False,
            skip_download=True,
        )
    except Exception as e:
        orig = (
            f" (orijinal: {type(first_exc).__name__}: {first_exc})"
            if first_exc else ""
        )
        raise RuntimeError(
            f"Drive klasör listeleme hatası: {type(e).__name__}: {e}{orig}. "
            "Klasör 'Anyone with the link' mi? ID/URL doğru mu?"
        )

    if not file_list:
        orig = f"Orijinal hata: {type(first_exc).__name__}: {first_exc}. " if first_exc else ""
        raise RuntimeError(
            f"{orig}Drive klasöründe indirilebilir dosya bulunamadı. "
            "Klasör 'Anyone with the link' erişimine açık mı?"
        )

    downloaded_paths: list[Path] = []
    skipped: list[str] = []

    # file_list: list[GoogleDriveFileToDownload(id, path, local_path)]
    for f in file_list:
        file_id: str = f.id
        local_path = Path(f.local_path)
        file_name: str = local_path.name

        # Zaten indirildiyse (birincil yol kısmen başarmış olabilir)
        if local_path.exists() and local_path.stat().st_size > 0:
            downloaded_paths.append(local_path)
            continue

        local_path.parent.mkdir(parents=True, exist_ok=True)

        # Native Google Docs/Sheets/Slides: klasör listesinde uzantısız isimle gelir.
        # gdown'ın kendi kaynak kodunda da bu not var: "Google-native files have no
        # extension in the folder listing."
        has_extension = bool(Path(file_name).suffix)
        if not has_extension:
            exported = _try_export_google_doc(file_id, local_path, quiet=quiet)
            if exported:
                downloaded_paths.append(exported)
            else:
                skipped.append(file_name or file_id)
                if not quiet:
                    print(
                        f"[bulk_import] Atlandı (native Google Doc, export başarısız): "
                        f"{file_name}",
                        file=_sys.stderr,
                    )
            continue

        # Normal dosya: önce requests ile dene (gdown 6.0.0'da
        # drive.usercontent.google.com redirect parse bug'ı var — curl/requests
        # ile aynı URL sorunsuz çalışıyor). Başarısız olursa gdown'a fallback.
        download_ok = _requests_download_drive_file(
            file_id, local_path, quiet=quiet
        )
        if download_ok and local_path.exists() and local_path.stat().st_size > 0:
            downloaded_paths.append(local_path)
            continue

        # Fallback: gdown.download (eski yol, bazı dosyalar için çalışabilir)
        try:
            result = gdown.download(
                url=f"https://drive.google.com/uc?id={file_id}",
                output=str(local_path),
                quiet=quiet,
                use_cookies=False,
            )
            result_path = Path(result) if result else local_path
            if result_path.exists() and result_path.stat().st_size > 0:
                downloaded_paths.append(result_path)
            elif local_path.exists() and local_path.stat().st_size > 0:
                downloaded_paths.append(local_path)
            else:
                skipped.append(file_name)
                if not quiet:
                    print(
                        f"[bulk_import] İndirme başarısız (boş sonuç): {file_name}",
                        file=_sys.stderr,
                    )
        except Exception as e:
            skipped.append(file_name)
            if not quiet:
                print(
                    f"[bulk_import] İndirme başarısız, atlandı: {file_name} "
                    f"→ {type(e).__name__}: {e}",
                    file=_sys.stderr,
                )

    if not downloaded_paths:
        orig = (
            f"Orijinal hata: {type(first_exc).__name__}: {first_exc}. "
            if first_exc else ""
        )
        skip_note = (
            f" Atlanan {len(skipped)} dosya: {', '.join(skipped[:5])}."
            if skipped else ""
        )
        raise RuntimeError(
            f"{orig}Drive klasöründen hiç dosya indirilemedi.{skip_note} "
            "Klasör 'Anyone with the link' erişimine açık mı?"
        )

    return downloaded_paths


def _requests_download_drive_file(
    file_id: str,
    dest: Path,
    quiet: bool = True,
) -> bool:
    """Public Drive dosyasını requests ile direkt indir.

    gdown 6.0.0'ın ``drive.usercontent.google.com`` redirect parse bug'ını
    bypass eder. curl'ün yaptığını Python'da yapar: 303 redirect'i takip et,
    içeriği yaz.

    Büyük dosyalar için Google "virüs taraması yapılamadı" uyarısı gösterebilir;
    ``confirm=t`` parametresi bunu otomatik geçer.

    Returns True on success (dosya yazıldı ve boş değil), False yoksa.
    """
    try:
        import requests as _req
    except ImportError:
        return False

    import sys as _sys

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        ),
    }
    url = f"https://drive.google.com/uc?id={file_id}&export=download&confirm=t"

    try:
        with _req.get(
            url,
            headers=headers,
            stream=True,
            allow_redirects=True,
            timeout=120,
        ) as resp:
            if resp.status_code != 200:
                if not quiet:
                    print(
                        f"[bulk_import] requests indirme: HTTP {resp.status_code} "
                        f"(file_id={file_id})",
                        file=_sys.stderr,
                    )
                return False

            # Content-Type kontrolü: HTML dönüyorsa muhtemelen onay sayfası
            ctype = resp.headers.get("content-type", "")
            if "text/html" in ctype.lower():
                # HTML onay sayfası geldi → virüs tarama uyarısı, parse edip
                # gerçek download URL'i bul (basit heuristic, çoğu dosyada
                # confirm=t parametresi yeterli olur, bu nadir bir kenar durum)
                if not quiet:
                    print(
                        f"[bulk_import] requests: HTML yanıt geldi "
                        f"(virus scan warning?) — atlanıyor (file_id={file_id})",
                        file=_sys.stderr,
                    )
                return False

            dest.parent.mkdir(parents=True, exist_ok=True)
            total = 0
            with open(dest, "wb") as f:
                for chunk in resp.iter_content(chunk_size=64 * 1024):
                    if chunk:
                        f.write(chunk)
                        total += len(chunk)

            if total == 0:
                if not quiet:
                    print(
                        f"[bulk_import] requests indirme: boş dosya "
                        f"(file_id={file_id})",
                        file=_sys.stderr,
                    )
                try:
                    dest.unlink()
                except OSError:
                    pass
                return False

            if not quiet:
                print(
                    f"[bulk_import] ✓ requests ile indirildi: {dest.name} "
                    f"({total} bytes)",
                    file=_sys.stderr,
                )
            return True

    except Exception as e:
        if not quiet:
            print(
                f"[bulk_import] requests indirme hatası "
                f"(file_id={file_id}): {type(e).__name__}: {e}",
                file=_sys.stderr,
            )
        return False


def _try_export_google_doc(
    file_id: str,
    dest_hint: Path,
    quiet: bool = True,
) -> Optional[Path]:
    """Native Google Doc'u export URL'i ile .docx veya .txt olarak indir.

    dest_hint: uzantısız hedef path. Başarılı exportta uygun uzantı eklenir.
    requests kütüphanesi yoksa None döner (gdown zaten bağımlılık olarak içeriyor).
    """
    try:
        import requests as _req
    except ImportError:
        return None

    import sys as _sys

    # Sırayla dene: önce Google Docs (document), sonra Sheets (spreadsheet)
    export_attempts = [
        ("document", "docx"),
        ("document", "txt"),
        ("spreadsheets", "csv"),
    ]

    for doc_type, fmt in export_attempts:
        export_url = (
            f"https://docs.google.com/{doc_type}/d/{file_id}/export?format={fmt}"
        )
        try:
            resp = _req.get(export_url, timeout=30, allow_redirects=True)
            if resp.status_code == 200 and len(resp.content) > 100:
                export_dest = dest_hint.with_suffix(f".{fmt}")
                export_dest.parent.mkdir(parents=True, exist_ok=True)
                export_dest.write_bytes(resp.content)
                if not quiet:
                    print(
                        f"[bulk_import] Native Google Doc export edildi "
                        f"({doc_type}/{fmt}): {export_dest.name}",
                        file=_sys.stderr,
                    )
                return export_dest
        except Exception:
            continue

    return None


def list_drive_folder_docx(folder_id_or_url: str,
                            tmp_dir: Optional[Path] = None) -> list[Path]:
    """Klasördeki desteklenen tüm dosyaların (docx/txt/md) lokal path'lerini döner.

    İsim 'docx' içeriyor ama backward-compat için bırakıldı —
    artık SUPPORTED_EXTENSIONS'ı tarar.

    Önizleme + actual import aynı download'u kullanır — gdown idempotent değil,
    o yüzden tmp_dir hem önizleme hem import için aynı kullanılmalı.
    """
    if tmp_dir is None:
        tmp_dir = Path(tempfile.mkdtemp(prefix="bulk_drive_"))
    files = download_drive_folder(folder_id_or_url, tmp_dir)
    return [p for p in files if p.suffix.lower() in SUPPORTED_EXTENSIONS]


# ---------------------------------------------------------------------------
# Bulk job creation
# ---------------------------------------------------------------------------
def pair_docx_with_lo(docx_paths: list[Path]) -> list[tuple[Path, Optional[Path]]]:
    """Drive'daki dokümanları main + _lo companion olarak grupla.

    Eşleştirme: 'senaryo1.<ext>' (main) + 'senaryo1_lo.<ext>' (learning objectives).
    Main ve LO **farklı uzantıda da olabilir** ('senaryo1.docx' + 'senaryo1_lo.txt'
    veya 'senaryo1.txt' + 'senaryo1_lo.docx' gibi). `_lo` suffix'iyle biten
    dosyalar companion olarak kabul edilir; ana dosya yoksa standalone (lo skip).

    Returns: [(main_path, lo_path_or_None), ...]
    Sıralama: main dosyaların stem'ine göre alfabetik.
    """
    by_stem: dict[str, Path] = {}
    for p in docx_paths:
        if p.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        # Aynı stem'de docx + txt varsa: docx'i öncele (daha zengin metadata)
        existing = by_stem.get(p.stem)
        if existing and existing.suffix.lower() == ".docx":
            continue
        by_stem[p.stem] = p

    pairs: list[tuple[Path, Optional[Path]]] = []
    seen: set[str] = set()
    for stem in sorted(by_stem.keys()):
        if stem in seen:
            continue
        if stem.endswith("_lo"):
            # Bu LO dosyası — ana script aynı klasörde varsa main loop'ta zaten
            # eşleştirilmiş olur. Burada gelirse main yok demektir, skip.
            main_stem = stem[:-3]
            if main_stem in by_stem:
                continue  # main loop tarafından zaten ele alındı
            # Standalone _lo, main yok → skip (warn)
            continue
        # Ana script
        lo_stem = stem + "_lo"
        lo_path = by_stem.get(lo_stem)
        pairs.append((by_stem[stem], lo_path))
        seen.add(stem)
        if lo_path is not None:
            seen.add(lo_stem)
    return pairs


def bulk_create_jobs_from_docx_paths(
    docx_paths: list[Path],
    *,
    custom_prompt_template: str = "",
    submitted_by: str = "bulk_import",
    job_factory: Callable[..., dict],  # caller'ın Job dataclass'ı için
    on_progress: Optional[Callable[[str], None]] = None,
) -> dict:
    """Her docx için bir Job dict oluşturur. job_factory(title, text, custom_prompt,
    submitted_by, learning_objectives) → dict döner (caller dataclass'a maple).

    `<name>_lo.docx` companion'ları otomatik eşleştirilir → ana job'ın
    `learning_objectives` field'ına geçirilir. Eşleşmeyen `_lo.docx` skip.

    Returns: {"created": [...job_dicts...], "errors": [(filename, err), ...]}
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(f"python-docx yüklü değil: {_DOCX_IMP_ERR}")

    created: list[dict] = []
    errors: list[tuple[str, str]] = []

    pairs = pair_docx_with_lo(docx_paths)
    if on_progress:
        n_paired = sum(1 for _, lo in pairs if lo is not None)
        on_progress(
            f"{len(pairs)} ana script → {n_paired} tanesi _lo.docx companion ile eşleşti."
        )

    for i, (main_path, lo_path) in enumerate(pairs):
        if on_progress:
            lo_note = f" + {lo_path.name}" if lo_path else ""
            on_progress(f"[{i+1}/{len(pairs)}] {main_path.name}{lo_note} işleniyor…")
        try:
            text = parse_document(main_path)
            if not text or len(text.strip()) < 50:
                errors.append((main_path.name, f"Çok kısa veya boş ({len(text)} chars)"))
                continue
            lo_text = ""
            if lo_path:
                try:
                    lo_text = parse_document(lo_path)
                except Exception as e:
                    errors.append((lo_path.name, f"LO parse fail: {e}"))
                    lo_text = ""
            title = docx_to_title(main_path.name, text)
            job = job_factory(
                title=title,
                text=text,
                custom_prompt=custom_prompt_template,
                submitted_by=submitted_by,
                learning_objectives=lo_text,
            )
            created.append(job)
        except Exception as e:
            errors.append((main_path.name, f"{type(e).__name__}: {str(e)[:200]}"))

    if on_progress:
        on_progress(
            f"Bitti: {len(created)} job oluşturuldu, {len(errors)} hatalı."
        )
    return {"created": created, "errors": errors}


# ---------------------------------------------------------------------------
# Combined: end-to-end (caller için tek-shot helper)
# ---------------------------------------------------------------------------
def bulk_import_from_drive(
    drive_url_or_id: str,
    custom_prompt_template: str,
    submitted_by: str,
    job_factory: Callable[..., dict],
    *,
    on_progress: Optional[Callable[[str], None]] = None,
    keep_downloads: bool = False,
) -> dict:
    """Tek çağrıda Drive klasör → docx download → Job dicts.

    Returns: {
        "created": [job_dicts],
        "errors": [(filename, err), ...],
        "downloads_dir": str (keep_downloads=True ise),
        "total_files": int (klasördeki toplam docx sayısı),
    }
    """
    ok, msg = is_available()
    if not ok:
        raise RuntimeError(msg)

    out_dir = Path(tempfile.mkdtemp(prefix="bulk_drive_"))
    if on_progress:
        on_progress(f"Drive klasörü indiriliyor → {out_dir}")
    try:
        docx_paths = list_drive_folder_docx(drive_url_or_id, out_dir)
        if on_progress:
            on_progress(f"{len(docx_paths)} adet .docx bulundu, işleniyor…")
        if not docx_paths:
            return {
                "created": [],
                "errors": [(drive_url_or_id, "Klasörde .docx dosyası yok veya erişim yok.")],
                "total_files": 0,
            }
        result = bulk_create_jobs_from_docx_paths(
            docx_paths,
            custom_prompt_template=custom_prompt_template,
            submitted_by=submitted_by,
            job_factory=job_factory,
            on_progress=on_progress,
        )
        result["total_files"] = len(docx_paths)
        if keep_downloads:
            result["downloads_dir"] = str(out_dir)
        else:
            try:
                shutil.rmtree(out_dir)
            except OSError:
                pass
        return result
    except Exception:
        if not keep_downloads:
            try:
                shutil.rmtree(out_dir, ignore_errors=True)
            except OSError:
                pass
        raise


__all__ = [
    "SUPPORTED_EXTENSIONS",
    "extract_folder_id",
    "is_available",
    "parse_docx",
    "parse_txt",
    "parse_document",
    "get_docx_metadata",
    "docx_to_title",
    "download_drive_folder",
    "list_drive_folder_docx",
    "pair_docx_with_lo",
    "bulk_create_jobs_from_docx_paths",
    "bulk_import_from_drive",
]
